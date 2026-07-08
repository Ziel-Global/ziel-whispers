from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

doc = Document()

# ── Styles ─────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── Helper Functions ───────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def new_page(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Ziel Log System')
run.font.size = Pt(42)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = tagline.add_run('Workforce Management Platform')
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Date: {date.today().strftime("%B %d, %Y")}\n').font.size = Pt(12)
meta.add_run('Version: 1.0\n').font.size = Pt(12)
meta.add_run('Technical Specification Document').font.size = Pt(12)

new_page(doc)

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. System Overview',
    '2. Admin Side — Full Feature List',
    '3. Employee Side — Full Feature List',
    '4. Feature Relationships & Dependencies',
    '5. Architecture Diagram',
    '6. Data Model — Current Database Tables',
    '7. Permissions Matrix',
    '8. Edge Functions & Scheduled Jobs',
]
for item in toc_items:
    doc.add_paragraph(item, style='Normal')

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. System Overview', level=1)

doc.add_paragraph(
    'Ziel Log System is a comprehensive workforce management SaaS platform designed to help '
    'organizations track employee attendance, work logs, leave requests, project assignments, '
    'and performance goals all in one place. The system replaces manual spreadsheets and email-based '
    'approval workflows with a centralized digital platform where administrators can manage their '
    'entire workforce while employees can self-serve their daily work tracking, leave applications, '
    'and attendance records.'
)

doc.add_paragraph(
    'Built with modern web technologies, the platform supports two primary roles: Admin and Employee '
    '(referred to as "Resource" within the system). Admins have full access to manage employees, '
    'projects, goals, reports, and system configuration. Employees can clock in/out, submit daily work '
    'logs, apply for leave, view their assigned projects and tasks, and manage their profile.'
)

doc.add_paragraph(
    'The technology stack consists of a React and TypeScript frontend deployed on Vercel, with a '
    'Supabase backend providing PostgreSQL database, authentication, Row-Level Security, Edge Functions '
    '(serverless TypeScript), and pg_cron for scheduled jobs. Email notifications are powered by the '
    'Resend API. All time calculations are locked to Pakistan Standard Time (PKT, UTC+5).'
)

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 2. ADMIN SIDE
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. Admin Side — Full Feature List', level=1)

# 2.1 Dashboard
doc.add_heading('2.1 Dashboard', level=2)
doc.add_paragraph(
    'The admin dashboard provides a centralized overview of the organization. It displays key statistics '
    'such as total active employees, present/absent counts for today, late arrivals, and pending leave '
    'and WFH requests requiring approval. A prominent "Team Today" banner shows the clock-in status of '
    'all employees at a glance. A "Late Employees" banner highlights employees who clocked in past the '
    'grace period. The dashboard also shows recent urgent announcements and provides quick-action links '
    'to common admin tasks.'
)

# 2.2 Employee Management
doc.add_heading('2.2 Employee Management', level=2)
doc.add_paragraph(
    'Admin can create, view, edit, and deactivate (soft-delete) employee accounts. The employee list '
    'supports filtering by status (active/inactive), employment type (full-time/part-time/contractor), '
    'role (admin/employee/client), and department. Each employee profile contains the following sections:'
)
bullets = [
    'Personal Information: Name, email, phone, department, designation, employment type, avatar upload.',
    'Shift Settings: Custom shift start/end times (overrides global defaults), working days (5 or 6), '
    'night shift flag, and overtime enable toggle.',
    'Log Edit Days: Per-employee override of how many past working days they can edit/submit logs for.',
    'Access Controls Tab: Remote Access toggle with optional from/to date range, Mark as On Leave with '
    'from/to date range, and Remote Site Access toggle (controls work mode selection).',
    'Leave Balances: View and manage per-year leave balances for each leave type (Sick, Casual, Hourly, Others).',
    'Monthly Logged Hours Summary: Bar showing expected vs actual logged hours for the current month.',
    'Oversight Flag: Toggle is_oversight for visual highlighting in admin views.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Bulk employee creation is supported via CSV import. Admins can also set/reset passwords, send '
    'invitation emails, and trigger password reset flows.'
)

# 2.3 Attendance Management
doc.add_heading('2.3 Attendance Management', level=2)
doc.add_paragraph(
    'The admin attendance view displays all clock-in/out records for a selected date. Records can be '
    'filtered by status (late/on time), work mode (onsite/remote), and employee. Key features include:'
)
bullets = [
    'Edit clock-in and clock-out times for any employee (with late recalculation triggered automatically '
    'by the database).',
    'CSV export of attendance data.',
    'Late/present/absent status indicators with color coding.',
    'Multi-session support — employees can clock in/out multiple times per day.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 2.4 Daily Logs Management
doc.add_heading('2.4 Daily Logs Management', level=2)
doc.add_paragraph(
    'Admin can view all submitted daily logs for any selected date, organized by employee. Each employee '
    'row shows logged hours, overtime, late submission status, and standup completion status. Features include:'
)
bullets = [
    'View per-employee log details with project name, category, hours, and description.',
    'Flag or lock logs for review (admin_flagged, admin_locked columns).',
    'Add a log entry on behalf of any employee (AdminAddLogDialog) with overtime awareness.',
    'Standup completion toggle per employee per day.',
    'Overtime hours breakdown (regular vs overtime hours).',
    'CSV export of daily logs.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 2.5 Leave Management
doc.add_heading('2.5 Leave Management', level=2)
doc.add_paragraph(
    'Full leave lifecycle management including request approval, rejection, balance tracking, and '
    'calendar visualization. The leave year runs June 1 to May 31.'
)
bullets = [
    'Leave Request Approval/Rejection: View all pending leave requests with employee details, leave type, '
    'date range, days count, and reason. Approve or reject with one click. Approval auto-updates leave '
    'balances and sends email notifications.',
    'Leave Balances: View and edit per-employee leave balances for each leave type per leave year. '
    'Supports adding custom balance rows.',
    'Leave Types: Sick Leave (paid, 10 days/year), Casual Leave (paid, 10 days/year), Hourly Leave '
    '(paid, tracked in hours), Others (unpaid).',
    'Hourly Leave: Employees can request partial-day leave in hours. Admin approval checks combined '
    'logged + leave hours against the daily expected hours threshold.',
    'Calendar View: Monthly calendar showing employee leave status for each day, with weekend blurring '
    'and leave-type color coding.',
    'Year System: Leave years (June-May) with automatic rollover display.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 2.6 WFH / Remote Work
doc.add_heading('2.6 Work From Home / Remote Work Requests', level=2)
doc.add_paragraph(
    'Admins can approve or reject employee WFH requests from a dedicated section. The WFH table shows '
    'requested date range, days count, reason, and current status. Expanded details show the full date '
    'range and working day count. Bulk remote access management allows enabling/disabling remote access '
    'for all non-admin employees at once with date range support. A scheduled cron job automatically '
    'clean up expired bulk remote access entries at midnight.'
)

# 2.7 Projects
doc.add_heading('2.7 Projects', level=2)
doc.add_paragraph(
    'Full project lifecycle management. Admin can create projects with name, description, client assignment, '
    'status, date range, and an optional document link. Each project detail page provides:'
)
bullets = [
    'Members Tab: Assign employees to the project with specific roles. Supports CSV import for bulk member '
    'assignment. Shows attendance pie chart for the project.',
    'Tasks Tab: Displays all tasks linked to the project, organized into three sections: Unlinked (tasks '
    'not assigned to any active goal), Active Goals (tasks grouped by goal), and Completed. Supports bulk '
    'task addition via CSV import.',
    'Phases Tab: Create and manage project phases with title, due date, status (unlinked/linked/'
    'in_progress/complete/returned), and sort order. Each phase has a dedicated edit page.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 2.8 Project Goals / Weekly Goals
doc.add_heading('2.8 Project Goals / Weekly Goals', level=2)
doc.add_paragraph(
    'Admin can create goals linked to projects with due dates and description. Goals are assigned to '
    'specific resources (employees) and each resource can have multiple tasks within the goal. The goal '
    'detail page shows all assigned resources with their tasks, task statuses (pending/in_progress/'
    'complete/returned), and overall goal progress calculated from task completion. Goals support status '
    'tracking (active/completed/cancelled) and task reassignment between resources.'
)

# 2.9 Clients
doc.add_heading('2.9 Clients', level=2)
doc.add_paragraph(
    'Client management with full CRUD operations. Each client record includes company name, contact '
    'person details, email, phone, address, industry classification, status (active/inactive/archived), '
    'and notes. Clients can be linked to projects and are filterable by industry and status.'
)

# 2.10 Reports
doc.add_heading('2.10 Reports', level=2)
doc.add_paragraph(
    'A comprehensive reporting suite with multiple report types, all exportable to CSV or PNG (heatmap):'
)
bullets = [
    'Utilization Report: Bar chart showing employee utilization percentage based on logged hours vs '
    'available working hours, with Low (<70%) / Good (70-110%) / Over (>110%) status indicators. '
    'Supports date range filtering and CSV export.',
    'Attendance Heatmap: Color-coded calendar heatmap per employee showing daily logged hours intensity '
    '(0h=muted, 1-6h=green, 6-9h=yellow, >9h=red). Supports month and department filtering. '
    'Exportable as PNG.',
    'Hourly Summary / Monthly Breakdown: Per-employee monthly summary showing total hours, late log '
    'count, project-hour breakdown, leave-by-type breakdown, and attendance mode distribution '
    '(onsite vs remote days).',
    'Attendance Trends: Line chart showing daily attendance rate (% of employees present) and average '
    'clock-in time over a date range. Filters by department.',
    'Daily Logs Report: Complete list of all submitted logs within a date range with employee, project, '
    'and description filters.',
    'Leave Report: All leave requests within a date range with employee and leave type details.',
    'Missed Logs Report: All detected missed log entries within a date range, filtered to exclude '
    'weekends (respecting per-employee working days), pre-account-creation dates, and approved leave days.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 2.11 Announcements
doc.add_heading('2.11 Announcements', level=2)
doc.add_paragraph(
    'Company announcement system with CRUD operations. Announcements support priority levels '
    '(normal/urgent), scheduled publishing (publish_at date), and audience targeting '
    '(all employees / specific departments). Urgent announcements appear as a dismissible banner '
    'on the employee dashboard. Read tracking per user shows who has viewed each announcement. '
    'The sidebar badge displays unread announcement count.'
)

# 2.12 Settings
doc.add_heading('2.12 Settings', level=2)
doc.add_paragraph(
    'System-wide configuration stored in the system_settings key-value table. Configurable settings include:'
)
bullets = [
    'Default shift start and end times.',
    'Expected daily hours (default 8).',
    'Late grace minutes (default 15).',
    'Annual leave entitlement in days.',
    'Timezone (locked to Asia/Karachi).',
    'Session timeout hours and max session lifetime (12h default).',
    'Auto clock-out time and display time.',
    'Log reminder offset minutes.',
    'Log edit window days (how many past working days employees can edit logs).',
    'Utilization thresholds (low/high percentages).',
    'Admin notification email address.',
    'Lockout window minutes and max failed login attempts.',
    'Menu ordering configuration.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 2.13 Audit Log
doc.add_heading('2.13 Audit Log', level=2)
doc.add_paragraph(
    'Complete audit trail of all system actions. Each audit record captures the actor (user who performed '
    'the action), action type (descriptive string like "user.created", "leave.approved"), target entity '
    'and ID, and optional metadata JSON. The audit log viewer supports filtering by action type and actor, '
    'with a detail modal showing full metadata. Actions can be exported to CSV. Tracked actions include '
    'employee CRUD, login/logout, client CRUD, project CRUD, log submission/editing/flagging/locking, '
    'attendance clock-in/out, leave and WFH approval/rejection, announcement CRUD, remote access bulk '
    'operations, and password changes.'
)

# 2.14 Email Notifications
doc.add_heading('2.14 Email Notifications', level=2)
doc.add_paragraph(
    'Automated email notifications are sent via the Resend API through Supabase Edge Functions. '
    'Notifications are triggered for the following events:'
)
bullets = [
    'New leave request: Notifies all admins with request details.',
    'New WFH request: Notifies all admins with request details.',
    'Leave request approved/rejected: Notifies the requesting employee.',
    'WFH request approved/rejected: Notifies the requesting employee.',
    'Log reminder: Automated email when shift end minus reminder offset is reached (cron-driven).',
    'Invitation email: Sent to new users with a "Set Your Password" link.',
    'All notifications are logged in the notifications table with delivery status and retry tracking.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 3. EMPLOYEE SIDE
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. Employee Side — Full Feature List', level=1)

# 3.1 Dashboard
doc.add_heading('3.1 Dashboard', level=2)
doc.add_paragraph(
    'Personal dashboard showing the employee\'s attendance status for today (clocked in/out, time elapsed), '
    'work mode (onsite/remote), late status with minutes, log submission status (submitted/not submitted), '
    'and a weekly log summary. Also displays a "Team Today" section showing which colleagues are clocked in, '
    'recent urgent announcements, and quick links to common actions.'
)

# 3.2 Clock In/Out
doc.add_heading('3.2 Clock In/Out', level=2)
doc.add_paragraph(
    'Employees can clock in and out from the attendance page with work mode selection (onsite or remote). '
    'The system enforces that clock-out time must be after clock-in. If remote access is not enabled for '
    'the employee, only "onsite" mode is available. When an employee submits daily logs and has an open '
    'session, they can optionally auto-clock-out at the same time. Late clock-ins are automatically '
    'calculated by a database trigger comparing clock-in time to shift start plus grace period. Multiple '
    'clock-in/out sessions are supported per day. An auto-clockout edge function cleans up stale sessions '
    'from previous dates at a configurable time.'
)

# 3.3 Submit Log
doc.add_heading('3.3 Submit Log', level=2)
doc.add_paragraph(
    'Daily work log submission with a comprehensive form. Employees select a project (from their assigned '
    'projects plus an automatic "Miscellaneous" option), select a task (from tasks assigned to them under '
    'the selected project), choose a category (development, meeting, bug_fix, code_review, deployment, '
    'documentation, testing, marketing, seo, research, posting, designing, outbound_calls, or other), '
    'enter hours (0.25–24), and provide a description (minimum 20 characters). The form includes:'
)
bullets = [
    'Daily progress bar showing logged vs expected hours.',
    'Date picker with weekend blocking (respects per-employee working days) and log edit window enforcement.',
    'Overtime marking when enabled — weekend hours and any hours beyond daily expected are flagged as overtime.',
    'Bulk submission: Multiple draft logs can be submitted at once.',
    'Past-date log submission within the configured log edit window (calculated in working days).',
    'Category constraint enforcement from database.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3.4 My Logs
doc.add_heading('3.4 My Logs', level=2)
doc.add_paragraph(
    'View submitted logs with filtering by date range and project. Each log entry shows date, project, '
    'category, hours, description, late flag, and overtime flag. Employees can edit or delete logs within '
    'the allowed edit window.'
)

# 3.5 My Attendance
doc.add_heading('3.5 My Attendance', level=2)
doc.add_paragraph(
    'Self-service attendance page featuring a monthly calendar view with color-coded attendance indicators '
    '(present/late/absent). Clicking a day shows detailed clock-in/out times, work mode, late minutes, and '
    'a live elapsed-time counter for the current session. The page calculates and displays a log reminder '
    'time based on shift end minus configured offset.'
)

# 3.6 Leave & Requests
doc.add_heading('3.6 Leave & Requests', level=2)
doc.add_paragraph(
    'Self-service leave application and WFH request system. Features include:'
)
bullets = [
    'Leave application with type selection (Sick, Casual, Hourly, Others), date range, reason, and '
    'optional document attachment.',
    'Duplicate validation: overlapping leave/WFH requests with pending or approved status are rejected.',
    'Weekend validation respects per-employee working days (5-day or 6-day work week).',
    'Hourly leave: Employees select hours (1 to max configurable) for partial-day leave.',
    'Leave balance display showing remaining days for each leave type.',
    'WFH requests: Date range selection with working day count calculation and reason field.',
    'Separate tabs for leave history and WFH request history.',
    'Annual history view showing all requests in the current leave year (June–May).',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3.7 My Projects
doc.add_heading('3.7 My Projects', level=2)
doc.add_paragraph(
    'View all assigned projects with project details. Each project shows a "Tasks" tab displaying tasks '
    'assigned to the employee under that project, grouped by goal with completion checkboxes. Task progress '
    'indicators show completion status. Employees can mark tasks as complete directly from this view.'
)

# 3.8 Announcements
doc.add_heading('3.8 Announcements', level=2)
doc.add_paragraph(
    'Read-only view of company announcements with pagination and read tracking. Urgent announcements '
    'appear as a dismissible banner on the dashboard. The sidebar shows an unread count badge.'
)

# 3.9 Profile
doc.add_heading('3.9 Profile', level=2)
doc.add_paragraph(
    'Read-only self-service profile view showing personal information, shift schedule, and account details. '
    'Employees can view their assigned working days, shift timings, and employment type.'
)

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 4. FEATURE RELATIONSHIPS & DEPENDENCIES
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. Feature Relationships & Dependencies', level=1)

doc.add_heading('4.1 Clock-In → Attendance, Late Detection, and Log Reminders', level=2)
doc.add_paragraph(
    'When an employee clocks in, an attendance record is created. A database trigger '
    '(calculate_late_clockin) automatically compares the clock-in time against the employee\'s shift '
    'start time plus grace period. If late, the is_late flag and minutes_late are set. The same trigger '
    'recalculates when an admin edits the clock-in time. The attendance record also stores a '
    'log_reminder_time (shift end minus reminder offset), which the send-log-reminder cron job uses to '
    'send email reminders if logs haven\'t been submitted. The "Team Today" section on both admin and '
    'employee dashboards reads today\'s attendance records to display who is clocked in.'
)

doc.add_heading('4.2 Log Submission → Tasks, Goals, and Phases', level=2)
doc.add_paragraph(
    'When an employee submits a daily log with an associated task (selected from assigned tasks), the log '
    'is stored in daily_logs. The task completion is not automatically tied to logs — employees must '
    'manually mark tasks as complete via the "My Projects" page. However, task completion status drives '
    'goal progress (percentage of completed tasks per goal) which is displayed on the goal detail page. '
    'Goal progress in turn feeds into phase progress for projects that use phases. The task-goal-phase '
    'hierarchy allows admins to track work at multiple levels of granularity.'
)

doc.add_heading('4.3 Leave Approval → Balances, Notifications, and On-Leave Status', level=2)
doc.add_paragraph(
    'When an admin approves a leave request, several things happen simultaneously: (1) The leave_balances '
    'table is updated to increment used_days for that employee, leave type, and leave year. (2) An email '
    'notification is sent to the employee via the send-request-notification edge function. (3) The '
    'attendance system treats the employee as on leave (affecting absent/late calculations). (4) The '
    'audit_log records the approval action. A rejected request sends a rejection email without modifying '
    'balances. Hourly leave approvals deduct from the hourly leave balance instead of full-day counts.'
)

doc.add_heading('4.4 Tasks → Goals → Phases', level=2)
doc.add_paragraph(
    'Tasks are created within goals, and goals are linked to projects. Each task has an assigned '
    'resource (employee), priority, description, and status (pending/in_progress/complete/returned). '
    'A goal\'s progress is calculated as the percentage of completed tasks out of total tasks. '
    'Projects can have phases, and tasks can optionally be linked to a phase. The Phase Edit page '
    'shows unlinked, linked, in-progress, completed, and returned tasks. This chain enables '
    'bottom-up progress tracking from individual tasks to project phases.'
)

doc.add_heading('4.5 Log Edit Days → Submit Log', level=2)
doc.add_paragraph(
    'The log_edit_days setting (either global or per-employee) controls how many past working days an '
    'employee can edit or submit logs for. The system calculates working days backward from today, '
    'excluding weekends based on the employee\'s working_days setting (5 or 6 days per week). Only days '
    'within this working-day window are enabled in the date picker. This prevents employees from modifying '
    'or back-dating logs beyond the allowed window.'
)

doc.add_heading('4.6 Access Controls → Clock-In Work Mode', level=2)
doc.add_paragraph(
    'The Access Controls section in Employee Profile provides three toggles that affect the employee\'s '
    'clock-in experience: (1) Remote Access — when enabled with a date range, the employee can select '
    '"remote" work mode during clock-in within that range. (2) Mark as On Leave — when active, the '
    'employee\'s attendance shows leave status. (3) Remote Site Access — when enabled, additional work '
    'mode options become available. The bulk remote access feature in LeaveAdmin allows admins to enable '
    'or disable remote access for all non-admin employees at once, with automatic cleanup of expired '
    'entries via a daily cron job.'
)

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 5. ARCHITECTURE DIAGRAM
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. Architecture Diagram', level=1)

doc.add_paragraph(
    'The following diagram shows how the main components of the system connect:'
)

lines = [
    '',
    '  ┌─────────────────────────────────────────────────────────────┐',
    '  │                    VERCEL (Frontend)                        │',
    '  │  React + TypeScript + Vite + Tailwind CSS + shadcn/ui      │',
    '  │  Pages: Dashboard, LogSubmit, MyLeave, Projects, etc.      │',
    '  └──────────┬──────────────────────────────────────┬──────────┘',
    '             │                              ▲',
    '             │ HTTPS (Supabase Client)       │ Server Actions',
    '             ▼                              │',
    '  ┌─────────────────────────────────────────────────────────────┐',
    '  │                   SUPABASE (Backend)                        │',
    '  │  ┌─────────────┐  ┌──────────────┐  ┌──────────────────┐   │',
    '  │  │ PostgreSQL  │  │  Auth        │  │  Edge Functions  │   │',
    '  │  │ Database    │  │  (User Mgmt) │  │  (Deno/TypeScript)│   │',
    '  │  │             │  │              │  │                  │   │',
    '  │  │ • 25 tables │  │ • JWT tokens │  │ • send-email     │   │',
    '  │  │ • RLS       │  │ • Sessions   │  │ • invite-user    │   │',
    '  │  │ • Triggers  │  │ • MFA-ready  │  │ • manage-user    │   │',
    '  │  │ • pg_cron   │  │              │  │ • send-invite    │   │',
    '  │  └─────────────┘  └──────────────┘  │ • log-login-att- │   │',
    '  │                                      │   empt            │   │',
    '  │                                      │ • send-log-      │   │',
    '  │                                      │   reminder        │   │',
    '  │                                      │ • send-request-   │   │',
    '  │                                      │   notification    │   │',
    '  │                                      │ • auto-clockout   │   │',
    '  │                                      │ • detect-missed-  │   │',
    '  │                                      │   logs            │   │',
    '  │                                      └────────┬─────────┘   │',
    '  └─────────────────────────────────────────────────────────────┘',
    '                                      │',
    '                                      ▼',
    '  ┌─────────────────────────────────────────────────────────────┐',
    '  │                    RESEND (Email Service)                    │',
    '  │  Sends: invitation emails, leave/WFH notifications,         │',
    '  │  log reminders. All delivery logged to notifications table. │',
    '  └─────────────────────────────────────────────────────────────┘',
    '',
    '  ┌─────────────────────────────────────────────────────────────┐',
    '  │              pg_cron SCHEDULED JOBS (in PostgreSQL)         │',
    '  │                                                             │',
    '  │  • Every 5 min:  send-log-reminder                         │',
    '  │  • Every 1 min:  auto-clockout (stale sessions)            │',
    '  │  • Scheduled:    detect-missed-logs (past shift end)        │',
    '  │  • Daily midnight: cleanup-expired-remote-access            │',
    '  └─────────────────────────────────────────────────────────────┘',
]

for line in lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 6. DATA MODEL
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. Data Model — Current Database Tables', level=1)

tables = [
    {
        'name': 'users',
        'purpose': 'Core user/employee profiles with role, shift, and access control flags.',
        'columns': [
            ('id', 'uuid PK', 'Primary key, default gen_random_uuid()'),
            ('email', 'text', 'Login email, unique'),
            ('password_hash', 'text', 'Hashed password'),
            ('full_name', 'text', 'Display name'),
            ('phone', 'text', 'Contact number'),
            ('department', 'text', 'Department name'),
            ('designation', 'text', 'Job title'),
            ('employment_type', 'text', 'full_time / part_time / contractor'),
            ('status', 'text', 'active / inactive'),
            ('role', 'text', 'admin / employee / client'),
            ('avatar_url', 'text', 'Profile image URL'),
            ('shift_start', 'time', 'Custom shift start time (nullable)'),
            ('shift_end', 'time', 'Custom shift end time (nullable)'),
            ('has_custom_shift', 'boolean', 'If true, use custom shift times instead of global defaults'),
            ('working_days', 'smallint', '5 (Mon-Fri) or 6 (Mon-Sat)'),
            ('is_night_shift', 'boolean', 'Night shift flag for late calculation rollover'),
            ('overtime_enabled', 'boolean', 'If true, user can log overtime hours'),
            ('log_edit_days', 'integer', 'Override global log edit window (nullable)'),
            ('reminder_offset_minutes', 'integer', 'Override global reminder offset (nullable)'),
            ('is_oversight', 'boolean', 'Visual highlight flag for admin views'),
            ('remote_access', 'boolean', 'Remote access enabled flag'),
            ('remote_access_from', 'date', 'Remote access start date (nullable)'),
            ('remote_access_to', 'date', 'Remote access end date (nullable)'),
            ('remote_access_bulk', 'boolean', 'True if set by bulk operation (nullable)'),
            ('remote_site_access', 'boolean', 'Additional remote site access flag'),
            ('is_on_leave', 'boolean', 'Currently on leave flag'),
            ('on_leave_from', 'date', 'Leave start date (nullable)'),
            ('on_leave_to', 'date', 'Leave end date (nullable)'),
            ('created_at', 'timestamptz', 'Account creation timestamp, not null'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'None (referenced by attendance, daily_logs, leave_requests, etc.)'
    },
    {
        'name': 'attendance',
        'purpose': 'Employee clock-in/out records with late tracking and work mode.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('date', 'date', 'Attendance date (PKT)'),
            ('clock_in', 'timestamptz', 'Clock-in timestamp'),
            ('clock_out', 'timestamptz', 'Clock-out timestamp (nullable)'),
            ('work_mode', 'text', 'onsite / remote'),
            ('is_late', 'boolean', 'Calculated by trigger'),
            ('minutes_late', 'integer', 'Late minutes, calculated by trigger'),
            ('hours_late', 'integer', 'Late hours, calculated by trigger'),
            ('log_reminder_time', 'timestamptz', 'When to send log reminder (computed client-side)'),
            ('note', 'text', 'Admin note (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'user_id → users.id'
    },
    {
        'name': 'daily_logs',
        'purpose': 'Daily work log entries submitted by employees.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee who logged'),
            ('project_id', 'uuid FK→projects', 'Project (nullable, shows as Miscellaneous)'),
            ('task_id', 'uuid FK→tasks', 'Associated task (nullable)'),
            ('log_date', 'date', 'Date the work was done'),
            ('category', 'text', 'Log category (development, meeting, bug_fix, etc.)'),
            ('hours', 'numeric', 'Hours worked (0.25–24)'),
            ('description', 'text', 'Work description (min 20 chars)'),
            ('status', 'text', 'draft / submitted'),
            ('is_late', 'boolean', 'Late submission flag (calculated)'),
            ('is_overtime', 'boolean', 'Overtime flag (calculated client-side)'),
            ('is_standup_done', 'boolean', 'Standup completion flag'),
            ('admin_flagged', 'boolean', 'Flagged by admin for review'),
            ('admin_locked', 'boolean', 'Locked by admin (prevent edits)'),
            ('submitted_at', 'timestamptz', 'Submission timestamp'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'user_id → users.id, project_id → projects.id, task_id → tasks.id'
    },
    {
        'name': 'daily_standups',
        'purpose': 'Per-user per-date standup completion tracking.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('date', 'date', 'Standup date'),
            ('is_done', 'boolean', 'Standup completed flag'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'user_id → users.id'
    },
    {
        'name': 'leave_requests',
        'purpose': 'Leave and WFH requests with approval workflow.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('leave_type_id', 'uuid FK→leave_types', 'Leave type (nullable for WFH)'),
            ('start_date', 'date', 'Leave/WFH start date'),
            ('end_date', 'date', 'Leave/WFH end date'),
            ('days_count', 'integer', 'Number of working days'),
            ('hours', 'numeric', 'Hours for hourly leave (nullable)'),
            ('reason', 'text', 'Reason for request'),
            ('document_url', 'text', 'Supporting document (nullable)'),
            ('status', 'text', 'pending / approved / rejected / cancelled'),
            ('reviewed_by', 'uuid FK→users', 'Admin who reviewed (nullable)'),
            ('reviewed_at', 'timestamptz', 'Review timestamp (nullable)'),
            ('rejection_reason', 'text', 'Reason if rejected (nullable)'),
            ('is_wfh', 'boolean', 'True if this is a WFH request'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'user_id → users.id, leave_type_id → leave_types.id, reviewed_by → users.id'
    },
    {
        'name': 'leave_balances',
        'purpose': 'Per-user per-year leave balance tracking.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('leave_type_id', 'uuid FK→leave_types', 'Leave type'),
            ('leave_year', 'text', 'Leave year label (e.g. "2025-2026")'),
            ('total_days', 'numeric', 'Total allocated days'),
            ('used_days', 'numeric', 'Days used so far'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'user_id → users.id, leave_type_id → leave_types.id'
    },
    {
        'name': 'leave_types',
        'purpose': 'Leave type definitions.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('name', 'text', 'Sick, Casual, Hourly, Others'),
            ('is_paid', 'boolean', 'Paid leave flag'),
            ('days_per_year', 'numeric', 'Annual entitlement days (nullable)'),
            ('color', 'text', 'Display color hex code'),
            ('sort_order', 'integer', 'Display order'),
            ('requires_hours', 'boolean', 'True for hourly leave type'),
        ],
        'fk': 'None'
    },
    {
        'name': 'remote_work_requests',
        'purpose': 'WFH request workflow table.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('start_date', 'date', 'WFH start'),
            ('end_date', 'date', 'WFH end'),
            ('days_count', 'integer', 'Working days in range'),
            ('reason', 'text', 'Reason'),
            ('status', 'text', 'pending / approved / rejected'),
            ('reviewed_by', 'uuid FK→users', 'Admin reviewer (nullable)'),
            ('reviewed_at', 'timestamptz', 'Review timestamp (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'user_id → users.id, reviewed_by → users.id'
    },
    {
        'name': 'projects',
        'purpose': 'Project definitions with client linkage.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('name', 'text', 'Project name'),
            ('description', 'text', 'Project description'),
            ('client_id', 'uuid FK→clients', 'Associated client (nullable)'),
            ('status', 'text', 'active / completed / archived'),
            ('start_date', 'date', 'Project start'),
            ('end_date', 'date', 'Project end (nullable)'),
            ('document_link', 'text', 'External document URL (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'client_id → clients.id'
    },
    {
        'name': 'project_members',
        'purpose': 'Employee-project assignments with roles.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('project_id', 'uuid FK→projects', 'Project'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('role', 'text', 'Role within project (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'project_id → projects.id, user_id → users.id'
    },
    {
        'name': 'project_roles',
        'purpose': 'Project role definitions.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('project_id', 'uuid FK→projects', 'Project'),
            ('name', 'text', 'Role name'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'project_id → projects.id'
    },
    {
        'name': 'project_phases',
        'purpose': 'Project phases with status tracking.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('project_id', 'uuid FK→projects', 'Project'),
            ('title', 'text', 'Phase title'),
            ('due_date', 'date', 'Due date (nullable)'),
            ('status', 'text', 'unlinked / linked / in_progress / complete / returned'),
            ('sort_order', 'integer', 'Display order'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'None (referenced by tasks)'
    },
    {
        'name': 'clients',
        'purpose': 'External client organizations linked to projects.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('name', 'text', 'Company name'),
            ('contact_person', 'text', 'Primary contact name'),
            ('email', 'text', 'Contact email'),
            ('phone', 'text', 'Contact phone'),
            ('address', 'text', 'Company address'),
            ('industry', 'text', 'Industry classification'),
            ('status', 'text', 'active / inactive / archived'),
            ('notes', 'text', 'Internal notes'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'None (referenced by projects.client_id)'
    },
    {
        'name': 'goals',
        'purpose': 'Project goals with progress tracking.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('project_id', 'uuid FK→projects', 'Associated project'),
            ('title', 'text', 'Goal title'),
            ('description', 'text', 'Goal description'),
            ('due_date', 'date', 'Due date (nullable)'),
            ('status', 'text', 'active / completed / cancelled'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'project_id → projects.id'
    },
    {
        'name': 'goal_resources',
        'purpose': 'Employee assignments to goals.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('goal_id', 'uuid FK→goals', 'Goal'),
            ('user_id', 'uuid FK→users', 'Assigned employee'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'goal_id → goals.id, user_id → users.id'
    },
    {
        'name': 'tasks',
        'purpose': 'Individual tasks within goals, assigned to resources.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('goal_id', 'uuid FK→goals', 'Parent goal'),
            ('project_id', 'uuid FK→projects', 'Associated project (nullable)'),
            ('phase_id', 'uuid FK→project_phases', 'Associated phase (nullable)'),
            ('assigned_to', 'uuid FK→users', 'Assigned employee'),
            ('title', 'text', 'Task title'),
            ('description', 'text', 'Task description'),
            ('priority', 'text', 'high / medium / low'),
            ('status', 'text', 'pending / in_progress / complete / returned'),
            ('due_date', 'date', 'Due date (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'goal_id → goals.id, project_id → projects.id, phase_id → project_phases.id, assigned_to → users.id'
    },
    {
        'name': 'announcements',
        'purpose': 'Company announcements with priority and scheduling.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('title', 'text', 'Announcement title'),
            ('content', 'text', 'Announcement body'),
            ('priority', 'text', 'normal / urgent'),
            ('audience', 'text', 'all / department-specific'),
            ('publish_at', 'timestamptz', 'Scheduled publish time (nullable)'),
            ('created_by', 'uuid FK→users', 'Admin who created it'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'created_by → users.id'
    },
    {
        'name': 'announcement_reads',
        'purpose': 'Per-user read/dismissed tracking for announcements.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('announcement_id', 'uuid FK→announcements', 'Announcement'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('read_at', 'timestamptz', 'When user read it'),
            ('dismissed', 'boolean', 'Dismissed flag (nullable)'),
        ],
        'fk': 'announcement_id → announcements.id, user_id → users.id'
    },
    {
        'name': 'notifications',
        'purpose': 'Email notification log with delivery status.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Recipient'),
            ('type', 'text', 'Notification type (leave, wfh, log_reminder, etc.)'),
            ('subject', 'text', 'Email subject'),
            ('body', 'text', 'Email body'),
            ('status', 'text', 'sent / failed / pending'),
            ('retry_count', 'integer', 'Number of delivery retries'),
            ('sent_at', 'timestamptz', 'Sent timestamp (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'user_id → users.id'
    },
    {
        'name': 'missed_logs',
        'purpose': 'Detected missed log entries (inserted by cron edge function).',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('log_date', 'date', 'Date of missed log'),
            ('reason', 'text', 'Reason for missing (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'user_id → users.id'
    },
    {
        'name': 'system_settings',
        'purpose': 'Key-value configuration store for system-wide settings.',
        'columns': [
            ('key', 'text PK', 'Setting key'),
            ('value', 'text', 'Setting value'),
            ('updated_by', 'uuid FK→users', 'Admin who last updated (nullable)'),
            ('updated_at', 'timestamptz', 'Last update timestamp'),
        ],
        'fk': 'updated_by → users.id'
    },
    {
        'name': 'audit_logs',
        'purpose': 'Complete audit trail for all entity actions.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('actor_id', 'uuid FK→users', 'User who performed the action'),
            ('action', 'text', 'Action description (user.created, leave.approved, etc.)'),
            ('target_entity', 'text', 'Entity type affected'),
            ('target_id', 'uuid', 'Entity ID (nullable)'),
            ('metadata', 'jsonb', 'Additional context data (nullable)'),
            ('created_at', 'timestamptz', 'Creation timestamp'),
        ],
        'fk': 'actor_id → users.id'
    },
    {
        'name': 'login_attempts',
        'purpose': 'Login attempt records for lockout detection.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('email', 'text', 'Attempted email'),
            ('ip_address', 'text', 'IP address'),
            ('success', 'boolean', 'Login success/failure'),
            ('attempted_at', 'timestamptz', 'Attempt timestamp'),
        ],
        'fk': 'None'
    },
    {
        'name': 'auto_clockout_acks',
        'purpose': 'User acknowledgments of auto clock-out events.',
        'columns': [
            ('id', 'uuid PK', 'Primary key'),
            ('user_id', 'uuid FK→users', 'Employee'),
            ('session_id', 'uuid FK→attendance', 'Auto-clocked-out session'),
            ('created_at', 'timestamptz', 'Acknowledgment timestamp'),
        ],
        'fk': 'user_id → users.id, session_id → attendance.id'
    },
]

for t in tables:
    doc.add_heading(f'6.{tables.index(t)+1} {t["name"]}', level=3)
    doc.add_paragraph(f'Purpose: {t["purpose"]}')
    headers = ['Column', 'Type', 'Description']
    add_table(doc, headers, t['columns'], col_widths=[4, 3.5, 10])
    p = doc.add_paragraph()
    run = p.add_run(f'Foreign Keys: {t["fk"]}')
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 7. PERMISSIONS MATRIX
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Permissions Matrix', level=1)
doc.add_paragraph('The following table shows which actions each role can perform:')

perm_headers = ['Feature / Action', 'Admin', 'Employee']
perm_rows = [
    ('View Dashboard', 'Yes', 'Yes (Personal)'),
    ('View Team Today', 'Yes (All)', 'Yes (Colleagues)'),
    ('Clock In / Out', 'Yes (Self)', 'Yes'),
    ('Edit Own Clock-In', 'No', 'No'),
    ('Edit Any Clock-In', 'Yes', 'No'),
    ('Submit Daily Logs', 'Yes (Self)', 'Yes'),
    ('Edit Own Logs (within window)', 'Yes', 'Yes'),
    ('Edit Any Employee\'s Logs', 'Yes', 'No'),
    ('Admin Add Log (on behalf)', 'Yes', 'No'),
    ('Flag / Lock Logs', 'Yes', 'No'),
    ('View Daily Logs (All)', 'Yes', 'No'),
    ('View Own Logs', 'Yes', 'Yes'),
    ('Create Employees', 'Yes', 'No'),
    ('Edit Employees', 'Yes', 'No'),
    ('Deactivate Employees', 'Yes', 'No'),
    ('Delete Employees', 'Yes', 'No'),
    ('CSV Import Employees', 'Yes', 'No'),
    ('Send Invitation Email', 'Yes', 'No'),
    ('Set/Reset Password', 'Yes', 'No'),
    ('Apply for Leave', 'Yes', 'Yes'),
    ('Approve/Reject Leave', 'Yes', 'No'),
    ('Edit Leave Balances', 'Yes', 'No'),
    ('Cancel Own Leave', 'Yes (own)', 'Yes'),
    ('Apply for WFH', 'Yes', 'Yes'),
    ('Approve/Reject WFH', 'Yes', 'No'),
    ('Bulk Remote Access Mgmt', 'Yes', 'No'),
    ('Create Projects', 'Yes', 'No'),
    ('Edit Projects', 'Yes', 'No'),
    ('Assign Project Members', 'Yes', 'No'),
    ('Create Goals', 'Yes', 'No'),
    ('Assign Goal Resources', 'Yes', 'No'),
    ('Create Tasks', 'Yes', 'No'),
    ('Reassign Tasks', 'Yes', 'No'),
    ('Manage Clients', 'Yes', 'No'),
    ('View Projects', 'Yes (All)', 'Yes (Assigned)'),
    ('Mark Task Complete', 'Yes', 'Yes'),
    ('Create Announcements', 'Yes', 'No'),
    ('View Announcements', 'Yes', 'Yes'),
    ('View Reports', 'Yes (All)', 'Limited'),
    ('View Audit Log', 'Yes', 'No'),
    ('Edit System Settings', 'Yes', 'No'),
    ('Edit Own Profile', 'Yes', 'Limited'),
    ('Edit Any Employee Profile', 'Yes', 'No'),
    ('View Attendance (All)', 'Yes', 'No'),
    ('View Own Attendance', 'Yes', 'Yes'),
    ('CSV Export', 'Yes', 'No'),
]

add_table(doc, perm_headers, perm_rows, col_widths=[7, 3, 3])

new_page(doc)

# ═══════════════════════════════════════════════════════════
# 8. EDGE FUNCTIONS & SCHEDULED JOBS
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Edge Functions & Scheduled Jobs', level=1)

doc.add_heading('8.1 Edge Functions (Supabase)', level=2)

ef_headers = ['Function', 'Description']
ef_rows = [
    ('send-email',
     'Core email sending service via Resend API. Called by other edge functions. '
     'Logs delivery status to the notifications table with retry tracking (2 retries). '
     'Does not require JWT authentication (config.toml sets verify_jwt = false).'),
    ('invite-user',
     'Admin-only endpoint. Creates a new auth user in Supabase Auth and inserts a corresponding '
     'profile row in the users table with default shift settings, working days, and reminder offset.'),
    ('manage-user',
     'Admin-only endpoint. Supports deactivating/reactivating users, setting/changing passwords, '
     'updating email addresses, fully deleting users (with cascading cleanup), and toggling the '
     'oversight flag.'),
    ('send-invite',
     'Sends a "Set Your Password" invitation email via the send-email function. Called after '
     'invite-user to deliver the sign-up link.'),
    ('log-login-attempt',
     'Records login attempts (success/failure) into the login_attempts table. Used by the lockout '
     'system which checks max_failed_login_attempts and lockout_window_minutes settings.'),
    ('send-log-reminder',
     'Cron-triggered function. Queries attendance records with log_reminder_time in the past '
     'and no submitted logs for the day, then sends email reminders to those employees. Runs '
     'every 5 minutes via pg_cron.'),
    ('send-request-notification',
     'Sends email notifications for leave and WFH requests. On new requests, notifies all admins. '
     'On approval/rejection, notifies the requesting employee. Called from MyLeave.tsx and LeaveAdmin.tsx.'),
    ('send-shift-reminder',
     'Currently disabled. Would send shift-start reminders to employees based on their shift times.'),
    ('send-escalation-reminder',
     'Currently disabled. Would send escalation reminders for employees who repeatedly miss logs.'),
    ('send-project-assignment',
     'Currently disabled. Would email employees when assigned to a new project.'),
    ('auto-clockout',
     'Cron-triggered function. Automatically clocks out stale attendance sessions from previous '
     'dates at the configured auto-clockout time. Logs acknowledgment records and writes to audit_log. '
     'Runs every minute via pg_cron.'),
    ('detect-missed-logs',
     'Cron-triggered function. Scans active employees who are past their shift end time without '
     'any submitted logs for today. Respects weekends (per-employee working days), approved leaves, '
     'and account creation dates. Inserts records into the missed_logs table.'),
]
add_table(doc, ef_headers, ef_rows, col_widths=[5, 13])

doc.add_paragraph()
doc.add_heading('8.2 Scheduled Jobs (pg_cron)', level=2)

cron_headers = ['Job Name', 'Schedule', 'Action']
cron_rows = [
    ('send-log-reminder', 'Every 5 minutes', 'Calls send-log-reminder edge function via http_get'),
    ('auto-clockout', 'Every 1 minute', 'Calls auto-clockout edge function'),
    ('detect-missed-logs', 'Scheduled (cron expression)', 'Calls detect-missed-logs edge function'),
    ('cleanup-expired-remote-access', 'Daily at midnight (PKT)', 'SQL function that clears remote_access flags '
     'and remote_access_bulk for users whose remote_access_to < current date'),
]
add_table(doc, cron_headers, cron_rows, col_widths=[5, 4, 9])

doc.add_paragraph()
doc.add_heading('8.3 External Services', level=2)
bullets = [
    'Resend: Email delivery API used by the send-email edge function for all transactional emails '
    '(invitations, notifications, reminders).',
    'Vercel: Frontend hosting and deployment platform.',
    'Supabase: Backend-as-a-Service providing PostgreSQL, Auth, Edge Functions, and storage.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ── Save ───────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ZielLogSystem_TechSpec.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
